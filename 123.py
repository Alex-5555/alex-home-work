from docx import Document
from docx.shared import Inches
from pprint import pprint

'''
print('Start')
def aprint(*b):
    print(len(b))
########################

aprint('Hello', 'Alex', 'ggg', 'jjjj')

def snek(d):
    if d%2 == 0:
        return True
    else:
        return False
for u in range(1,300):
    print(snek(u))

def home(g):
    def flat():
        # print('yyyyyyyyyyy')
        return 'ok'
    return flat
# print(home()())
# f = home()
# print(f)
#f()
#print(f())
#####################
# @home
def Alex():
    return 'jjj' 
# g = Alex((home()))
# print(Alex(home()()))
# print(g())

1 год = 100 грн. + 100 грн./10
2 год = 110 грн. + 110 грн./10
3 

years = [2022, 2023, 2024]
koef = 5
depos = 900
for y in years:
    depos = depos + depos/koef

def bank (a, p, years):
    counter = 0
    while counter < years:
        counter = counter + 1
        a = a + a/10
    return a
print(bank(100, 3))

def ppp(deposit, years):
    counter = 0
    while counter < years:
        counter = counter + 1
        deposit = deposit + deposit/10
    return deposit
print(ppp(100, 5))

print((200/100)*10)
'''
import sys
# print(sys.argv[2])

def ppp():
    counter = 0
    deposit = int(sys.argv[1])
    years = int(sys.argv[2])
    procent = int(sys.argv[3])
    while counter < years:
        counter = counter + 1
        deposit = deposit + (deposit/100)*procent
    return deposit
print(ppp())

document = Document()

document.add_heading('Розрахунок', 0)

p = document.add_paragraph(str(ppp()))

document.save('demo.docx')

  